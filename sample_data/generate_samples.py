import os
import email
from email.message import EmailMessage
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib import colors
import docx

def create_sample_1_pdf():
    pdf_path = "sample_data/Sample_1_Foreign_Particles_Paracetamol.pdf"
    doc = SimpleDocTemplate(pdf_path, pagesize=letter, rightMargin=40, leftMargin=40, topMargin=40, bottomMargin=40)
    styles = getSampleStyleSheet()
    
    title_style = ParagraphStyle(
        'DocTitle',
        parent=styles['Heading1'],
        fontSize=18,
        leading=22,
        textColor=colors.HexColor("#1e3a8a")
    )
    
    body_style = ParagraphStyle(
        'DocBody',
        parent=styles['Normal'],
        fontSize=10,
        leading=14
    )
    
    elements = []
    elements.append(Paragraph("<b>GLOBAL PHARMACEUTICAL QUALITY ASSURANCE</b>", title_style))
    elements.append(Paragraph("<b>FORMAL CUSTOMER COMPLAINT NOTIFICATION</b>", styles['Heading3']))
    elements.append(Spacer(1, 15))
    
    meta_data = [
        ["Reporting Customer:", "Apex Pharma International Ltd.", "Date Reported:", "16 August 2026"],
        ["Product Name:", "Paracetamol API (Acetaminophen)", "Grade / Strength:", "USP Grade (99.5% Purity)"],
        ["Batch / Lot Number:", "PA240812", "Affected Quantity:", "25 kg (1 fiber drum)"],
        ["Manufacturing Date:", "12 August 2026", "Expiry Date:", "August 2028"],
        ["Complaint Type:", "Foreign Matter / Visible Particulate", "Initial Severity:", "High / Urgent Triage"]
    ]
    
    t = Table(meta_data, colWidths=[130, 170, 110, 120])
    t.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, -1), colors.HexColor("#f8fafc")),
        ('TEXTCOLOR', (0, 0), (-1, -1), colors.HexColor("#0f172a")),
        ('FONTNAME', (0, 0), (-1, -1), 'Helvetica'),
        ('FONTSIZE', (0, 0), (-1, -1), 9),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
        ('TOPPADDING', (0, 0), (-1, -1), 6),
        ('GRID', (0, 0), (-1, -1), 0.5, colors.HexColor("#cbd5e1")),
    ]))
    elements.append(t)
    elements.append(Spacer(1, 15))
    
    elements.append(Paragraph("<b>1. Detailed Incident Description:</b>", styles['Heading4']))
    desc = (
        "During standard incoming raw material sampling and dispensation at our formulation facility (Suite 4B), "
        "our quality control technicians observed multiple visible black particulate matter suspended within the top layer "
        "of the 25 kg fiber drum of Paracetamol API (Batch PA240812). "
        "The particles appear to be 0.5 mm to 1.2 mm in size with metallic/elastomeric characteristics. "
        "Immediate quarantine was placed on drum serial # DRUM-004. No material from this batch has been charged into production. "
        "Retention samples and photographic evidence are available upon request."
    )
    elements.append(Paragraph(desc, body_style))
    elements.append(Spacer(1, 10))
    
    elements.append(Paragraph("<b>2. Requested Quality Actions:</b>", styles['Heading4']))
    actions = (
        "1. Immediate investigation into manufacturing line filtration and centrifuge integrity.<br/>"
        "2. Provision of a Certificate of Analysis (CoA) and root-cause investigation report.<br/>"
        "3. Authorization for return of the compromised 25 kg drum and expedited replacement delivery."
    )
    elements.append(Paragraph(actions, body_style))
    elements.append(Spacer(1, 20))
    
    elements.append(Paragraph("<i>Signed: Dr. Elena Vance, Senior QA Director — Apex Pharma International</i>", body_style))
    doc.build(elements)
    print(f"Created {pdf_path}")

def create_sample_2_txt():
    txt_path = "sample_data/Sample_2_Packaging_Defect_Amoxicillin.txt"
    content = """CUSTOMER COMPLAINT TRANSMISSION
=====================================================
TO: Quality Assurance & Regulatory Compliance Team
FROM: BioHealth Distribution Corp
DATE: 14 August 2026
SUBJECT: Packaging Defect Notification - Amoxicillin Trihydrate 500mg Capsules

Customer Details:
- Customer Name: BioHealth Distribution Corp
- Receiving Facility: Central Logistics Hub 3, Dallas TX

Product & Batch Details:
- Product Name: Amoxicillin Trihydrate Capsules
- Strength: 500mg
- Batch / Lot Number: AMX-2026-884
- Manufacturing Date: 10 May 2026
- Expiry Date: May 2029
- Quantity Received: 5,000 cartons
- Quantity Affected: 350 cartons (compromised secondary packaging)

Incident Description:
Upon receiving shipping container #CONT-88219, our warehouse inspectors identified that 350 outer shipping cartons of Amoxicillin Trihydrate 500mg (Batch AMX-2026-884) exhibited crushed corners and compromised tamper-evident security tape. Preliminary inspection showed the inner primary blister foils remained intact without moisture ingress, but the secondary packaging fails commercial GDP presentation standards and cannot be distributed to hospital pharmacies.

Complaint Classification: Packaging / Labeling Defect
Severity: Medium
Priority: Normal

Please provide return merchandise authorization (RMA) and replacement packaging units.
"""
    with open(txt_path, "w", encoding="utf-8") as f:
        f.write(content)
    print(f"Created {txt_path}")

def create_sample_3_docx():
    docx_path = "sample_data/Sample_3_Incorrect_Strength_Ibuprofen.docx"
    doc = docx.Document()
    
    doc.add_heading("URGENT: Product Quality Query & Out-of-Specification Report", 0)
    
    p = doc.add_paragraph()
    p.add_run("Customer Name: ").bold = True
    p.add_run("Nordic Care Pharmaceuticals AB\n")
    p.add_run("Complaint Date: ").bold = True
    p.add_run("15 August 2026\n")
    p.add_run("Product Name: ").bold = True
    p.add_run("Ibuprofen Granules DC\n")
    p.add_run("Product Strength / Grade: ").bold = True
    p.add_run("85% Direct Compression Grade (Target: 400mg equivalent)\n")
    p.add_run("Batch / Lot Number: ").bold = True
    p.add_run("IBU-DC-9011\n")
    p.add_run("Manufacturing Date: ").bold = True
    p.add_run("01 June 2026\n")
    p.add_run("Expiry Date: ").bold = True
    p.add_run("June 2029\n")
    p.add_run("Quantity Affected: ").bold = True
    p.add_run("1,200 kg\n")
    
    doc.add_heading("Detailed Description of Deviation:", level=2)
    doc.add_paragraph(
        "Nordic Care QC analytical testing of incoming Ibuprofen DC Granules (Batch IBU-DC-9011) "
        "yielded potency assay results of 72.4%, failing the certified release specification of 85.0% - 87.0% DC grade. "
        "This potency discrepancy represents an Out-of-Specification (OOS) event that could result in sub-potent finished dosage tablets. "
        "All 1,200 kg from pallet lots 01 through 06 have been placed under strict QA quarantine in our cold-storage hold area. "
        "This issue is flagged as High Severity and requires immediate technical investigation."
    )
    
    doc.add_heading("Immediate Requirements:", level=2)
    doc.add_paragraph(
        "1. Complete investigation into blending uniformity and granulator parameters.\n"
        "2. Review of release Certificate of Analysis from manufacturing site.\n"
        "3. Teleconference with Head of Manufacturing Quality within 24 hours."
    )
    
    doc.save(docx_path)
    print(f"Created {docx_path}")

def create_sample_4_eml():
    eml_path = "sample_data/Sample_4_Email_Complaint_Metformin.eml"
    msg = EmailMessage()
    msg['Subject'] = "Customer Complaint - Discoloration in Metformin HCl 500mg (Batch MET-26-04)"
    msg['From'] = "dr.kavita.patel@sunhealthpharma.com"
    msg['To'] = "complaints.qa@aivoa-pharma.com"
    msg['Date'] = "Mon, 17 Aug 2026 09:15:00 +0000"
    
    body = """Dear AIVOA Quality Assurance Team,

We are filing an official product complaint regarding Metformin HCl 500mg Film-Coated Tablets, supplied under Batch Number MET-26-04 (Mfg Date: 15 July 2026, Expiry: July 2029).

Customer Name: SunHealth Pharma Ltd.
Quantity Affected: 15,000 blister packs (150 shipper cartons)

Description:
During routine stability and dispensary verification at our regional warehouse in Mumbai, multiple blister strips from Batch MET-26-04 were found to exhibit yellowish-brown surface discoloration and mottled tablet appearance instead of the standard uniform white film coat. 

Given the physical appearance alteration and potential degradation/oxidation indicators, we have locked the batch in our ERP system and initiated distribution hold.

Please log this complaint, initiate batch record review, and confirm sample collection protocol.

Best regards,
Dr. Kavita Patel
Head of Quality Operations
SunHealth Pharma Ltd.
"""
    msg.set_content(body)
    with open(eml_path, "wb") as f:
        f.write(msg.as_bytes())
    print(f"Created {eml_path}")

if __name__ == "__main__":
    os.makedirs("sample_data", exist_ok=True)
    create_sample_1_pdf()
    create_sample_2_txt()
    create_sample_3_docx()
    create_sample_4_eml()
