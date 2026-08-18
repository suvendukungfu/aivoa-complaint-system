import io
import email
from email import policy
from typing import List, Dict, Any, Tuple
import docx
from pypdf import PdfReader
from fastapi import UploadFile, HTTPException

from backend.app.core.config import settings

def extract_pages_from_pdf(content: bytes) -> List[Dict[str, Any]]:
    """Extract structured pages with 1-based page numbers from PDF byte content"""
    try:
        reader = PdfReader(io.BytesIO(content))
        pages = []
        for idx, page in enumerate(reader.pages):
            page_text = page.extract_text() or ""
            if page_text.strip():
                pages.append({
                    "page_number": idx + 1,
                    "text": page_text.strip()
                })
        
        if not pages:
            raise HTTPException(
                status_code=422,
                detail="This PDF document appears to contain no extractable text. OCR is not enabled in this demo."
            )
        return pages
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"Failed to parse PDF document: {str(e)}")

def extract_text_from_pdf(content: bytes) -> str:
    """Extract complete text from PDF byte content"""
    pages = extract_pages_from_pdf(content)
    return "\n\n".join(p["text"] for p in pages)

def extract_text_from_docx(content: bytes) -> str:
    """Extract text from DOCX byte content"""
    try:
        doc = docx.Document(io.BytesIO(content))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        
        # Also extract table text if present
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    paragraphs.append(row_text)
                    
        extracted = "\n".join(paragraphs).strip()
        if not extracted:
            raise HTTPException(status_code=422, detail="The uploaded Word document contains no readable text.")
        return extracted
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"Failed to parse DOCX document: {str(e)}")

def extract_text_from_txt(content: bytes) -> str:
    """Extract text from TXT byte content"""
    try:
        # Try UTF-8 first, fallback to Latin-1
        try:
            text = content.decode("utf-8")
        except UnicodeDecodeError:
            text = content.decode("latin-1")
            
        cleaned = text.strip()
        if not cleaned:
            raise HTTPException(status_code=422, detail="The text document is empty.")
        return cleaned
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"Failed to decode TXT document: {str(e)}")

def extract_text_from_eml(content: bytes) -> str:
    """Extract text from EML email content including metadata headers"""
    try:
        msg = email.message_from_bytes(content, policy=policy.default)
        headers = [
            f"Email Subject: {msg.get('Subject', 'No Subject')}",
            f"From: {msg.get('From', 'Unknown')}",
            f"To: {msg.get('To', 'Unknown')}",
            f"Date: {msg.get('Date', 'Unknown')}"
        ]
        
        body_parts = []
        if msg.is_multipart():
            for part in msg.walk():
                ctype = part.get_content_type()
                cdispo = str(part.get('Content-Disposition'))
                if ctype == "text/plain" and "attachment" not in cdispo:
                    body = part.get_content()
                    if body:
                        body_parts.append(body)
        else:
            body = msg.get_content()
            if body:
                body_parts.append(body)
                
        full_text = "\n".join(headers) + "\n\n" + "\n".join(body_parts)
        cleaned = full_text.strip()
        if not cleaned:
            raise HTTPException(status_code=422, detail="The email contains no readable body text.")
        return cleaned
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"Failed to parse EML email document: {str(e)}")

def extract_document_pages(content: bytes, ext: str) -> List[Dict[str, Any]]:
    """
    Extract structured document pages.
    PDF returns actual 1-based page numbers.
    DOCX, TXT, EML return page_number=None since they do not have pagination metadata.
    """
    ext_lower = ext.lower().strip()
    if ext_lower == ".pdf":
        return extract_pages_from_pdf(content)
    elif ext_lower == ".docx":
        text = extract_text_from_docx(content)
        return [{"page_number": None, "text": text}]
    elif ext_lower == ".txt":
        text = extract_text_from_txt(content)
        return [{"page_number": None, "text": text}]
    elif ext_lower == ".eml":
        text = extract_text_from_eml(content)
        return [{"page_number": None, "text": text}]
    else:
        raise HTTPException(status_code=400, detail=f"Unsupported format '{ext}'.")

async def parse_uploaded_file(file: UploadFile) -> Tuple[str, str, int]:
    """Validate and extract text from an uploaded document"""
    filename = file.filename or "unknown"
    ext = "." + filename.split(".")[-1].lower() if "." in filename else ""
    
    if ext not in settings.ALLOWED_EXTENSIONS:
        raise HTTPException(
            status_code=400,
            detail=f"Unsupported file format '{ext}'. Supported formats: {', '.join(settings.ALLOWED_EXTENSIONS)}"
        )
        
    content = await file.read()
    file_size = len(content)
    
    if file_size > settings.MAX_UPLOAD_SIZE_BYTES:
        raise HTTPException(
            status_code=413,
            detail=f"File exceeds maximum allowed limit of {settings.MAX_UPLOAD_SIZE_BYTES / (1024*1024):.0f}MB"
        )
        
    if ext == ".pdf":
        text = extract_text_from_pdf(content)
    elif ext == ".docx":
        text = extract_text_from_docx(content)
    elif ext == ".txt":
        text = extract_text_from_txt(content)
    elif ext == ".eml":
        text = extract_text_from_eml(content)
    else:
        raise HTTPException(status_code=400, detail="Unsupported file format.")
        
    return filename, text, file_size
